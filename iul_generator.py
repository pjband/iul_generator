#!/usr/bin/env python
# -*- coding: utf-8 -*-

import os
import datetime
import time
import hashlib
import docx
import copy
import shutil
import re
import traceback

lang = {
    'byte': 'Б',
    'file': 'Файл',
    'changed': 'Изменен',
    'size': 'Размер',
    'checksum_sha1': 'Контр. сумма SHA-1',
    'list_of_documents': 'Ведомость электронных документов',
    'identification_sheet': 'ИУЛ',
    'identification_sheet_title': 'ВЕДОМОСТЬ ЭЛЕКТРОННЫХ ДОКУМЕНТОВ',

    'directory': 'Директория',
    'wrong_directory': 'Несуществующая директория',
    'creator': 'Разработчик ИУЛ',
    'normcontroller': 'Нормоконтролер',
    'project_name': 'Номер проета',
    'result': 'Результат',
    'iul_msg': 'Не забудьте проверить правильность наименований ИУЛ в ИУЛ.docx'
}

os.environ['TZ'] = 'Europe/Moscow'
cfg = {}
cfg['exclude_types'] = ('.ini', '.log', '.err')
cfg['clearFile'] = 'clear.docx'
cfg['newIUL'] = lang['identification_sheet']+'.docx'
cfg['fileName'] = lang['list_of_documents']+'.txt'
cfg['fileTitle'] = lang['identification_sheet_title']+'\r\n\r\n'\
        + lang['file']+': '+cfg['fileName']+'\r\n   '\
        + lang['changed']+': '+str(time.strftime("%d.%m.%y %X", time.localtime()))+'\r\n\r\n'
# cfg['exclude'] = 'ИУЛ'


# crutch
def copy_table_after(table, paragraph):
    tbl, p = table._tbl, paragraph._p
    new_tbl = copy.deepcopy(tbl)
    p.addnext(new_tbl)


def add_table(template, repl, clearFile):
    iul = docx.Document(clearFile)
    tmpl = docx.Document(template)
    tpl = tmpl.tables[0]
    for row in tpl.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for key in repl.keys():
                    paragraph.text = paragraph.text.replace(str(key), str(repl[key]))
    # add table in the end
    copy_table_after(tpl, iul.paragraphs[-1])
    iul.add_paragraph('')
    # page break
    if len(iul.tables) % 2 == 0:
        iul.add_page_break()
    iul.save(clearFile)


def input_dir():
    r = {}
    r['dir'] = str(input(lang['directory'] + ': '))
    if os.path.exists(r['dir']) is False:
        print (lang['wrong_directory'])
        return False
    else:
        r['creator'] = str(input(lang['creator'] + ': '))
        r['normcontroller'] = str(input(lang['normcontroller'] + ': '))
        return r


def iul_name(filename):
    iul_name = re.search(r'\d{2}-\d{4}-\w{1,3}', filename)
    if iul_name is None:
        return filename
    else:
        return iul_name.group(0)


if __name__ == "__main__":

    while True:
        inp =  input_dir()
        if inp:
            root_dir = inp['dir']
            break

    repl = {}
    i = 1
    
    newIULFile = os.path.join(root_dir, cfg['newIUL'])
    try:
        shutil.copy(cfg['clearFile'], newIULFile)
    except:
        print(traceback.format_exc())

    # list of documents:
    with open(root_dir+'/'+cfg['fileName'], 'w', encoding='utf8') as f:
        f.write(cfg['fileTitle'])
        for top, dirs, nondirs in os.walk(root_dir):
            for name in nondirs:
                path = str(os.path.join(top, name))
                if path.lower().endswith(cfg['exclude_types']) is False and name != cfg['newIUL'] and name != cfg['fileName']:
                    filename = os.path.basename(path)
                    modTime = time.strftime("%d.%m.%y %X", time.localtime(os.stat(path).st_mtime))
                    size = os.stat(path).st_size
                    try:
                        sha1 = hashlib.sha1(open(path, 'rb').read()).hexdigest()
                        f.write(
                            lang['file'] + ': ' + str(filename) + '\r\n' +
                            '   ' + lang['changed'] + ': ' + str(modTime) + '\r\n' +
                            '   ' + lang['size']+': ' + str(size) + ' ' + lang['byte'] + '\r\n' +
                            '   ' + lang['checksum_sha1'] + ': ' + str(sha1) + '\r\n\r\n'
                        )
                        # changing identification sheet
                        repl['Npp'] = str(i)
                        i = i + 1
                        repl['file'] = filename
                        repl['sha-1_hash'] = sha1
                        repl['creator'] = inp['creator']
                        repl['normcontroller'] = inp['normcontroller']
                        repl['iul_name'] = iul_name(filename)
                        repl['date'] = time.strftime("%d.%m.%y", time.localtime())
                        repl['num'] = str(i // 2)
                        repl['count'] = ''
                        add_table('iul_tempale.docx', repl, newIULFile)
                    except:
                        print('ERROR!!! File: ' + path)

    print(lang['result'] + ': ' + cfg['fileName'])
    print(lang['iul_msg'])
    input('Enter')
    exit()
